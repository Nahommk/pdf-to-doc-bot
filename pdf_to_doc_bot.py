#!/usr/bin/env python3
"""
PDF to DOC Telegram Bot
Converts PDF files to Microsoft Word DOC format (97-2003)
"""

import os
import logging
import tempfile
import shutil
from pathlib import Path
from telegram import Update
from telegram.ext import Application, CommandHandler, MessageHandler, ContextTypes, filters
import subprocess
import platform

# For PDF text extraction
import PyPDF2
from pdfplumber import PDF
import python_docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Enable logging
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# Your Telegram Bot Token (Get from @BotFather)
BOT_TOKEN = os.environ.get('BOT_TOKEN')

if not BOT_TOKEN:
    raise ValueError("No BOT_TOKEN found in environment variables!")

# Maximum file size (20MB)
MAX_FILE_SIZE = 20 * 1024 * 1024

class PDFToDocConverter:
    """Handle PDF to DOC conversion"""
    
    @staticmethod
    def extract_text_from_pdf(pdf_path):
        """Extract text content from PDF"""
        text_content = []
        
        try:
            # Try with pdfplumber first (better formatting preservation)
            with PDF.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append({
                            'page': page_num,
                            'text': page_text,
                            'tables': page.extract_tables()
                        })
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")
            
            # Fallback to PyPDF2
            try:
                with open(pdf_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        text = page.extract_text()
                        text_content.append({
                            'page': page_num + 1,
                            'text': text,
                            'tables': []
                        })
            except Exception as e:
                logger.error(f"Failed to extract text from PDF: {e}")
                raise
                
        return text_content
    
    @staticmethod
    def create_doc_from_text(text_content, output_path):
        """Create DOC file from extracted text"""
        doc = Document()
        
        # Add document title
        doc.add_heading('Converted from PDF', 0)
        
        for page_data in text_content:
            # Add page separator
            if page_data['page'] > 1:
                doc.add_page_break()
            
            # Add page header
            page_header = doc.add_paragraph(f"Page {page_data['page']}")
            page_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            page_header.runs[0].font.bold = True
            page_header.runs[0].font.size = Pt(10)
            page_header.runs[0].font.color.rgb = RGBColor(128, 128, 128)
            
            # Add tables if present
            if page_data['tables']:
                for table_data in page_data['tables']:
                    if table_data and len(table_data) > 0:
                        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                        table.style = 'Table Grid'
                        
                        for i, row_data in enumerate(table_data):
                            for j, cell_data in enumerate(row_data):
                                if cell_data:
                                    table.rows[i].cells[j].text = str(cell_data)
                        
                        doc.add_paragraph()  # Add spacing after table
            
            # Add text content
            if page_data['text']:
                # Split text into paragraphs
                paragraphs = page_data['text'].split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        para = doc.add_paragraph(para_text.strip())
                        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Save as DOCX first
        docx_path = output_path.replace('.doc', '.docx')
        doc.save(docx_path)
        
        # Convert DOCX to DOC using LibreOffice (if available)
        if shutil.which('libreoffice') or shutil.which('soffice'):
            try:
                cmd = 'libreoffice' if shutil.which('libreoffice') else 'soffice'
                subprocess.run([
                    cmd, '--headless', '--convert-to', 'doc',
                    '--outdir', os.path.dirname(output_path),
                    docx_path
                ], check=True, capture_output=True)
                os.remove(docx_path)  # Remove temporary DOCX file
                return output_path
            except Exception as e:
                logger.warning(f"LibreOffice conversion failed: {e}")
        
        # If LibreOffice not available, rename DOCX to DOC
        # (Most modern Word versions can open it)
        os.rename(docx_path, output_path)
        return output_path
    
    @staticmethod
    def convert_pdf_to_doc(pdf_path, output_path):
        """Main conversion function"""
        try:
            # Extract text from PDF
            text_content = PDFToDocConverter.extract_text_from_pdf(pdf_path)
            
            if not text_content:
                raise ValueError("No text content could be extracted from PDF")
            
            # Create DOC from extracted text
            doc_path = PDFToDocConverter.create_doc_from_text(text_content, output_path)
            
            return doc_path
        except Exception as e:
            logger.error(f"Conversion failed: {e}")
            raise

async def start_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Send a welcome message when the command /start is issued."""
    welcome_text = """
📄 *PDF to DOC Converter Bot* 📄

Welcome! I can convert your PDF files to Microsoft Word DOC format.

✨ *Features:*
• Convert PDF to DOC (Word 97-2003 format)
• Extract text and tables
• Preserve formatting where possible
• Support for multi-page documents

📤 *How to use:*
1. Send me any PDF file
2. I'll process and convert it
3. Download your DOC file

⚡ *Commands:*
/start - Show this welcome message
/help - Get help and information
/stats - View conversion statistics

Just send me a PDF file to get started!
"""
    await update.message.reply_text(welcome_text, parse_mode='Markdown')

async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Send a help message when the command /help is issued."""
    help_text = """
🤖 *PDF to DOC Converter Help*

*Commands:*
• /start - Welcome message
• /help - This help message
• /stats - Your conversion statistics
• /about - About this bot

*How to Convert:*
1. Send a PDF file (up to 20MB)
2. Wait for processing
3. Receive your DOC file

*Supported Features:*
✅ Text extraction
✅ Table preservation
✅ Multi-page documents
✅ Unicode text support

*Limitations:*
• Max file size: 20MB
• Complex layouts may lose formatting
• Images are not preserved (text only)
• Scanned PDFs need OCR (not supported)

*Tips:*
• For best results, use text-based PDFs
• Complex formatting may be simplified
• Tables are preserved when possible

Need help? Contact @YourUsername
"""
    await update.message.reply_text(help_text, parse_mode='Markdown')

async def about_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Send information about the bot"""
    about_text = """
ℹ️ *About PDF to DOC Bot*

Version: 1.0.0
Developer: Your Name

This bot converts PDF files to Microsoft Word DOC format (97-2003).

*Technical Details:*
• Language: Python 3
• Libraries: PyPDF2, pdfplumber, python-docx
• Format: DOC (Word 97-2003)

*Privacy:*
• Files are deleted after conversion
• No data is stored permanently
• Secure processing

*Source Code:*
Available on GitHub

*Support:*
Contact @YourUsername for help
"""
    await update.message.reply_text(about_text, parse_mode='Markdown')

async def stats_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Show user statistics"""
    user_id = update.effective_user.id
    
    # Get user stats from context (or database in production)
    user_data = context.user_data
    conversions = user_data.get('conversions', 0)
    total_size = user_data.get('total_size', 0)
    
    stats_text = f"""
📊 *Your Statistics*

User ID: `{user_id}`
Total Conversions: *{conversions}*
Total Data Processed: *{total_size / (1024*1024):.2f} MB*

Thank you for using PDF to DOC Bot!
"""
    await update.message.reply_text(stats_text, parse_mode='Markdown')

async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle incoming documents"""
    try:
        document = update.message.document
        
        # Check if document exists
        if not document:
            await update.message.reply_text("❌ No document received. Please send a PDF file.")
            return
        
        # Check file extension
        file_name = document.file_name
        if not file_name.lower().endswith('.pdf'):
            await update.message.reply_text(
                "❌ Invalid file format!\n\n"
                "Please send a PDF file with .pdf extension.\n"
                f"You sent: {file_name}"
            )
            return
        
        # Check file size
        if document.file_size > MAX_FILE_SIZE:
            await update.message.reply_text(
                f"❌ File too large!\n\n"
                f"Maximum size: {MAX_FILE_SIZE / (1024*1024):.0f} MB\n"
                f"Your file: {document.file_size / (1024*1024):.2f} MB"
            )
            return
        
        # Send processing message
        processing_msg = await update.message.reply_text(
            "⏳ *Processing your PDF...*\n\n"
            f"📄 File: {file_name}\n"
            f"📊 Size: {document.file_size / 1024:.2f} KB\n\n"
            "Please wait...",
            parse_mode='Markdown'
        )
        
        # Download the file
        file = await document.get_file()
        
        # Create temporary directory
        with tempfile.TemporaryDirectory() as temp_dir:
            # Download PDF
            pdf_path = os.path.join(temp_dir, file_name)
            await file.download_to_drive(pdf_path)
            
            # Update processing message
            await processing_msg.edit_text(
                "⚙️ *Converting to DOC format...*\n\n"
                "This may take a moment for large files.",
                parse_mode='Markdown'
            )
            
            # Convert PDF to DOC
            output_filename = file_name.replace('.pdf', '.doc').replace('.PDF', '.doc')
            output_path = os.path.join(temp_dir, output_filename)
            
            try:
                PDFToDocConverter.convert_pdf_to_doc(pdf_path, output_path)
            except Exception as e:
                await processing_msg.edit_text(
                    f"❌ *Conversion failed!*\n\n"
                    f"Error: {str(e)}\n\n"
                    "This might be a scanned PDF or have complex formatting.\n"
                    "Try with a different PDF file.",
                    parse_mode='Markdown'
                )
                return
            
            # Check if output file was created
            if not os.path.exists(output_path):
                await processing_msg.edit_text(
                    "❌ Failed to create DOC file.\n"
                    "Please try again or contact support."
                )
                return
            
            # Send the converted file
            await processing_msg.edit_text(
                "📤 *Uploading your DOC file...*",
                parse_mode='Markdown'
            )
            
            with open(output_path, 'rb') as doc_file:
                await update.message.reply_document(
                    document=doc_file,
                    filename=output_filename,
                    caption=(
                        f"✅ *Conversion Complete!*\n\n"
                        f"📄 Original: {file_name}\n"
                        f"📝 Converted: {output_filename}\n"
                        f"📊 Size: {os.path.getsize(output_path) / 1024:.2f} KB\n\n"
                        "Enjoy your DOC file!"
                    ),
                    parse_mode='Markdown'
                )
            
            # Delete processing message
            await processing_msg.delete()
            
            # Update user statistics
            user_data = context.user_data
            user_data['conversions'] = user_data.get('conversions', 0) + 1
            user_data['total_size'] = user_data.get('total_size', 0) + document.file_size
            
            logger.info(f"Successfully converted {file_name} for user {update.effective_user.id}")
            
    except Exception as e:
        logger.error(f"Error handling document: {e}")
        await update.message.reply_text(
            "❌ An unexpected error occurred.\n"
            "Please try again later or contact support."
        )

async def handle_non_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle non-document messages"""
    await update.message.reply_text(
        "📄 Please send me a PDF file to convert.\n\n"
        "Use /help for more information."
    )

async def error_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Log errors and notify user"""
    logger.error(f"Update {update} caused error {context.error}")
    
    if update and update.effective_message:
        await update.effective_message.reply_text(
            "❌ An error occurred while processing your request.\n"
            "Please try again or contact support."
        )

def main():
    """Start the bot"""
    # Create application
    application = Application.builder().token(BOT_TOKEN).build()
    
    # Register handlers
    application.add_handler(CommandHandler("start", start_command))
    application.add_handler(CommandHandler("help", help_command))
    application.add_handler(CommandHandler("about", about_command))
    application.add_handler(CommandHandler("stats", stats_command))
    
    # Document handler
    application.add_handler(MessageHandler(filters.Document.PDF, handle_document))
    application.add_handler(MessageHandler(filters.Document.ALL & ~filters.Document.PDF, handle_document))
    
    # Text message handler
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_non_document))
    
    # Error handler
    application.add_error_handler(error_handler)
    
    # Start the bot
    print("🤖 PDF to DOC Bot is starting...")
    print("Press Ctrl+C to stop")
    
    application.run_polling(allowed_updates=Update.ALL_TYPES)

if __name__ == '__main__':
    main()
